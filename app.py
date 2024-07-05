import streamlit as st
from docx import Document
from datetime import datetime
import pdfkit
import tempfile

# Función para cargar el documento de plantilla
def load_template(file_path):
    return Document(file_path)

# Función para reemplazar los marcadores en el documento
def replace_markers(doc, replacements):
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)
    return doc

# Función para guardar el documento modificado
def save_document(doc, file_path):
    doc.save(file_path)

# Función para convertir documento a PDF
def convert_to_pdf(docx_path, pdf_path):
    pdfkit.from_file(docx_path, pdf_path)

# Crear la interfaz de Streamlit
st.title("Generador de Oferta de Préstamo")

# Ingresar datos
mes = st.text_input("Mes", value=datetime.now().strftime("%B"))
dia = st.text_input("Día", value=datetime.now().strftime("%d"))
cliente = st.text_input("Cliente")
interes = st.text_input("Interés")
prestamista = st.text_input("Prestamista")
comitente_prestamista = st.text_input("Comitente Prestamista")

# Botón para generar el documento
if st.button("Generar Documento"):
    # Ruta fija al documento de plantilla
    template_path = '/mnt/data/MODELO A - Cohen SA - Oferta de préstamo de VN - COHEN TOMADOR.docx'
    
    # Cargar el documento de plantilla
    doc = load_template(template_path)
    
    # Reemplazar los marcadores con los datos ingresados
    replacements = {
        "[MES]": mes,
        "[DIA]": dia,
        "[CLIENTE]": cliente,
        "[INTERES]": interes,
        "[PRESTAMISTA]": prestamista,
        "[COMITENTEPRESTAMISTA]": comitente_prestamista
    }
    
    doc = replace_markers(doc, replacements)
    
    # Guardar el documento modificado en un archivo temporal
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_docx:
        docx_path = tmp_docx.name
        save_document(doc, docx_path)
    
    # Convertir el documento a PDF en un archivo temporal
    with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_pdf:
        pdf_path = tmp_pdf.name
        convert_to_pdf(docx_path, pdf_path)
    
    # Mostrar éxito y botón de descarga
    st.success("Documento generado exitosamente.")
    st.download_button(
        label="Descargar Documento PDF",
        data=open(pdf_path, "rb").read(),
        file_name="Oferta_de_Préstamo_Completada.pdf",
        mime="application/pdf"
    )
